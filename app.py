import os
import re
import streamlit as st
from streamlit_option_menu import option_menu
import numpy as np
import pandas as pd
import xlsxwriter
import scipy
import plotly.express as px
import sklearn
import xgboost as xgb
from typing_extensions import Concatenate
import validators
import hashlib
import openai
import requests

from langchain.callbacks.streamlit import StreamlitCallbackHandler
from langchain_groq import ChatGroq

from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import PromptTemplate, ChatPromptTemplate, MessagesPlaceholder
from langchain.schema.runnable import RunnableLambda, Runnable
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

from langchain_community.document_loaders import WikipediaLoader, PyPDFLoader, YoutubeLoader, UnstructuredURLLoader
from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings

from langchain.vectorstores import FAISS
from langchain.schema import Document
from langchain.retrievers import BM25Retriever, MergerRetriever
from langchain.chains import RetrievalQA, create_retrieval_chain, create_history_aware_retriever
## Code #####

def get_text_from_audio(audio_file, groq_api_key):
    with open(audio_file, 'rb') as audio_fp:
        files = {'file': (audio_fp.name, audio_fp, 'audio/m4a')}
        data = {'model': 'whisper-large-v3','language': 'en', 'response_format': 'text'}
        headers = {'Authorization': f'Bearer {groq_api_key}'}
        endpoint = "https://api.groq.com/openai/v1/audio/transcriptions"
        response = requests.post(endpoint, headers=headers, files=files, data=data)
    try:
      result = response.json()
    except Exception as e:
      raise ValueError(f"Could not decode JSON: {e}\nRaw response: {response.text}")
    if response.status_code == 200:
      text = response.json()['text']
      if os.path.exists(audio_file):
        os.remove(audio_file)
    else:
      raise RuntimeError(f"Groq Whisper API error: {result}")
    if "text" not in result:
      raise KeyError(f"No 'text' in response: {result}")

    return text

code_assistant_prompt = """You are CodeGenie, an expert software engineer and coding tutor.
You are currently helping a user named {username}.
Your job is to help {username} with code suggestions, debugging, and explanations across programming languages like Python, Java, C++, JavaScript, SQL, etc.
Your reply style should be:
- Friendly and encouraging (start with phrases like "Great question!", "Sure!", or "Let's walk through it...")
- Clear, concise answers
- Relevant code blocks
- Helpful comments and explanations
- Address the user by name when appropriate
- Use the language asked by the user
- Keep extra text minimal, but don’t be robotic
"""
code_assistant_title = "🤖 Your Coding Assistant"
code_assistant_header = """
    <div style="padding: 10px; text-align: center;">
    <h2 style="color: #4CAF50; font-size: 28px; margin-bottom: 0;">💻 Your personal Coding Assistant!</h2>
    <p style="font-size: 16px;">
        Get instant help with code suggestions, debugging support, and clear explanations across a wide range of programming languages — including Python, Java, C++, JavaScript, SQL....  
        <br>
        Simplify your coding journey with AI-powered assistance! 🚀
    </p>
</div>
    """
math_assistant_prompt = """
You are an expert mathematics tutor helping a user named {username}.
Your job is to solve mathematical questions of all kinds, including arithmetic, algebra, geometry, calculus, statistics, linear algebra, and word problems.
Please follow these guidelines:
1. Break the problem down into clear, logical steps.
2. Show all work for calculations and justify any rules or theorems used.
3. If the problem is word-based, first extract the relevant information and formulate equations.
4. If multiple approaches exist, briefly mention them, and then choose the most efficient one.
5. If the final answer is numerical, round it to a reasonable number of decimal places if needed.
6. Use LaTeX formatting to display math expressions cleanly if supported.
Your reply style should be friendly and encouraging (start with phrases like "Great question!", "Sure!", or "Let's walk through it...")
Let’s begin solving the problem.
"""
math_assistant_title = "🤖 Your Math Assistant"
math_assistant_header = """
    <div style="padding: 10px; text-align: center;">
    <h2 style="color: #4CAF50; font-size: 28px; margin-bottom: 0;">🧮 Your personal Math Assistant!</h2>
    <p style="font-size: 16px;">
        Simply type your math question, and let the assistant guide you step-by-step through the solution.  
        <br>
        Make problem-solving easier and smarter with AI! 💡
    </p>
</div>
    """
text_summarization_prompt = """
You are an expert language assistant helping a user named {username}.
Your task is to read and summarize any text provided by the user — whether it's an article, email, report, research paper, or story.

Please follow these guidelines:
1. Understand the main idea and supporting details of the text.
2. Summarize the content clearly and concisely, keeping the most important points.
3. Maintain the original tone and intent (informative, persuasive, formal, etc.).
4. If the text is technical or academic, retain key terminology but simplify explanations where helpful.
5. Avoid including unnecessary examples or repetition from the original.
6. Offer summaries in bullet points if appropriate, or in a short paragraph for general content.
7. Keep your tone helpful and friendly (start with phrases like "Sure!", "Here's a quick summary:", etc.).
Let’s get started with the summarization.
"""
text_summarization_title = "🤖 Your Text Summarizer"
text_summarization_header = """
    <div style="padding: 10px; text-align: center;">
    <h2 style="color: #4CAF50; font-size: 28px; margin-bottom: 0;">📝 Your Text Summarizer!</h2>
    <p style="font-size: 16px;">
        Effortlessly condense articles, emails, reports, or any text into clear and concise summaries — all within seconds.  
        <br>
        Just paste your content and let AI do the rest! 💡
    </p>
</div>
    """
Excel_Analyser_prompt = """
You are a helpful and friendly data analyst proficient in data analysis, data visualization, data preprocessing, and machine learning, assisting a user named {username}.
The user has uploaded a data file, which is already loaded into a Pandas DataFrame named `df`.
Your job is to strictly follow every command given by the user, one at a time.
You MUST follow these rules and output format:
- Use only the provided DataFrame `df` in your solution. DO NOT use `pd.read_csv` or `pd.read_excel` under any condition.
- The uploaded data has the following structure:
    - Columns: {columns}
    - Sample rows: {head}

COMMAND GUIDELINES:
1. If the user asks to **display the result**:
    - Perform the requested operation.
    - Assign the final DataFrame to a variable named `result`.
    - When displaying an already computed variable, assign it correctly to 'result' without modifying the variable content.
    - NEVER introduce a typo in the variable name.
    - ALWAYS double-check and match the variable name exactly as created in previous steps.

2. If the user asks to **plot chart**:
    - Use only Plotly Express for plotting (`import plotly.express as px`).
    - Assign the resulting figure to a variable named `fig`.
    - Do NOT use `matplotlib`, `seaborn`, or any other libraries unless the user explicitly asks for it.
    - NEVER use `fig.show()` or `plt.show()`.

3. If the user asks to **use matplotlib**:
    - Only then can you use `matplotlib.pyplot` and related functions (like `plot_tree`).
    - Display-related functions like `plt.show()` are still NOT allowed.
    - Assign the resulting figure to a variable named `fig`.
    - Still, NEVER use `fig.show()` or `plt.show()`.

4. If the user asks to **export file**:
    - Perform the operation and create a new DataFrame if needed.
    - Export it using `.to_excel('output.xlsx', index=False)`.
    - Assign `'output.xlsx'` to a variable named `result`.

5. If a previously generated code resulted in an error and the user gives a follow-up command to fix it:
    - Identify and fix the issue.
    - Return a clean, corrected full code block.

RESPONSE FORMAT:
- Write only clean, executable code in a single code block.
- Always include clear and concise **inline comments** using `######` to explain each step.
- Do NOT include any markdown, plain text, or explanations outside the code.
- Do NOT use backticks or code fences.

ALWAYS follow the user's latest instructions accurately and without deviation.
"""

Excel_Analyser_title = "🤖 Your Data Analyzer Bot"
Excel_Analyser_header = """
    Welcome to your personal **Data Analyzer Bot**!
    Generate the code in seconds. 
    Just paste the query, and get code! 💡
    """
Rag_chatbot_prompt = """
You are an assistant for question-answering tasks. Use the following pieces of retrieved context to answer the question. If you don't know the answer,
say that you don't know. Use three sentences maximum and keep the answer concise.
{context}
"""
Rag_chatbot_title = "🤖 Your RAG-Based Chatbot"
Rag_chatbot_header = """
    Welcome to your personal **RAG-Based Chatbot**!
    You can ask the question based on your uploaded content.
    """
contextualize_q_system_prompt = (
    "Given a chat history and the latest user question "
    "which might reference context in the chat history, "
    "formulate a standalone question which can be understood "
    "without the chat history. Do NOT answer the question, "
    "just reformulate it if needed and otherwise return it as is.")

groq_api_key = st.secrets["GROQ_API_KEY"]
PASSWORD = st.secrets["PASSWORD"]

def get_prompt(tool, user_name):
    if tool == "💻 Code Assistant":
        title = code_assistant_title
        header = code_assistant_header
        assistant_content = f"Hi {user_name}, I'm a code assistant. How can I help you?"
        system_prompt = code_assistant_prompt
    elif tool == "🧮 Math Assistant":
        title = math_assistant_title
        header = math_assistant_header
        assistant_content = f"Hi {user_name}, I'm a math assistant. How can I help you?"
        system_prompt = math_assistant_prompt
    elif tool == "📝 Text Summarizer":
        title = text_summarization_title
        header = text_summarization_header
        assistant_content = f"Hi {user_name}, I'm a Text Summarizer. How can I help you?"
        system_prompt = text_summarization_prompt
    elif tool == "📊 Data Analyzer Bot":
        title = Excel_Analyser_title
        header = Excel_Analyser_header
        assistant_content = f"Hi {user_name}, I'm a Data Analyzer Bot. How can I help you?"
        system_prompt = Excel_Analyser_prompt
    elif tool == "🔎 RAG-based Chatbot":
        title = Rag_chatbot_title
        header = Rag_chatbot_header
        assistant_content = f"Hi {user_name}, I'm a RAG Based Chatbot. How can I help you?"
        system_prompt = Rag_chatbot_prompt
    output_dict = {"title":title, "header":header, "assistant_content":assistant_content, "system_prompt":system_prompt}
    return output_dict
    
def rag_chatbot_uploader():
    st.markdown("""
                <div style="padding: 10px; text-align: center;">
            <h2 style="color: #4CAF50; font-size: 28px; margin-bottom: 0;">
                🤖 Your RAG-Based Chatbot!
            </h2>
            <p style="font-size: 17px; margin-top: 5px;">
                Upload your documents to serve as retrievers for generating accurate, context-aware responses from the Language Model (LLM).
                <br>
                Empower your conversations with knowledge-driven AI! 🚀
            </p>
        </div>
        """, unsafe_allow_html=True)
    tab1, tab2, tab3 = st.tabs(["Upload a File", "Enter Website URL", "Enter Text Manually"])
    user_input = None
    # Center align using columns
    col1, col2, col3 = st.columns([1,2,1])
    with tab1:
        uploaded_file = st.file_uploader("Upload a pdf/docs/txt file", type=["pdf", "docx", "txt"])
        generate_file_input = st.button("📂 Load file")
        
    with tab2:
        url_input = st.text_input("Website, Wikipedia URL", placeholder="Enter the URL here...")
        generate_url_input = st.button("🚀 Go Ahead")
    
    with tab3:
        text_input = st.text_area(
            "Text Input", height=150, placeholder="Paste your text here..."
        )
        generate_text_input = st.button("👍 Go Ahead")
    
    user_input = None
    with col2:
        if generate_file_input and uploaded_file is not None:
            try:
                with st.spinner("🔄 Uploading..."):
                    file_type = uploaded_file.type
                    if file_type == "text/plain":
                        user_input = str(uploaded_file.read(), "utf-8")
                    elif file_type == "application/pdf":
                        pdfreader = PdfReader(uploaded_file)
                        user_input = ''
                        for i, page in enumerate(pdfreader.pages):
                            content = page.extract_text()
                            if content:
                                user_input += content
                    elif (file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                        doc = docx.Document(uploaded_file)
                        user_input = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                st.exception(f"Exception:{e}")
        elif generate_text_input and text_input:
            user_input = text_input
        
        elif generate_url_input and url_input:
            if not url_input.strip():
                st.error("Please provide the information to get started")
            elif not validators.url(url_input):
                st.error("Please enter a valid Url. It should be a Wiki/website url")
            else:
                try:
                    with st.spinner("🔄 Uploading..."):
                        if "en.wikipedia.org" in url_input:
                            query = url_input.split("/")[-1]
                            loader = WikipediaLoader(query=query, load_max_docs=2)
                        else:
                            loader=UnstructuredURLLoader(urls=[url_input],ssl_verify=False,
                                                         headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_5_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/116.0.0.0 Safari/537.36"})
                        docs=loader.load()
                        raw_texts = [doc.page_content for doc in docs]
                        clean_texts = [re.sub(r'\s*\n\s*', ' ', text) for text in raw_texts]
                        clean_texts = [re.sub(r'\s{2,}', ' ', text).strip() for text in clean_texts]
                        user_input = "\n".join([text for text in clean_texts])
                except Exception as e:
                    st.exception(f"Exception:{str(e)[0:500] + '.....'}")
    
        if user_input:
            with st.spinner("🔄 Retriever loading..."):
                embedding_model = HuggingFaceEmbeddings(model_name="BAAI/bge-base-en-v1.5")
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=800,
                    chunk_overlap=20,
                    separators=["\n\n", "\n", ".", "!", "?", " ", ""])
                documents = [Document(page_content=chunk) for chunk in splitter.split_text(user_input)]
                faiss_db = FAISS.from_documents(documents, embedding_model)
                dense_retriever = faiss_db.as_retriever(search_type="mmr",  # MMR = Maximal Marginal Relevance for relevance + diversity
                                                search_kwargs={"k": 3, # number of docs
                                                               "lambda_mult": 0.7}) # (1.0 = pure relevance, 0.0 = pure diversity)
                # BM25Retriever is a keyword-based retriever
                bm25_retriever = BM25Retriever.from_documents(documents)
                bm25_retriever.k = 3
                hybrid_retriever = MergerRetriever(retrievers=[dense_retriever, bm25_retriever],
                                           weights=[0.7, 0.3])
                st.session_state.context = user_input
                st.session_state.hybrid_retriever = hybrid_retriever
                st.session_state.retriever_ready = True
                st.success(f"Loaded successfully.")
                st.rerun()
        # else:
        #     st.error(f"Please provide your file/url/text.")

def data_analysis_uploader():
    st.markdown("""
    <div style="padding: 10px; text-align: center;">
    <h2 style="color: #4CAF50; font-size: 28px; margin-bottom: 0;">
        🤖 Your Data Analyzer Bot!
    </h2>
    <p style="font-size: 17px; margin-top: 5px;">
        Upload your data files, ask anything, and instantly get code, results, and visual charts.
        <br>
        Simplify learning, research, and data exploration with the power of AI! 🚀
    </p>
</div>
    """, unsafe_allow_html=True)

    if "data" not in st.session_state:
        st.session_state.data = None
        st.session_state.file_name = ""
    uploaded_file = st.file_uploader("Upload your Excel or CSV file", type=["xlsx", "xls", "csv"])
    if uploaded_file and st.button("Load File for Analysis"):
        try:
            with st.spinner("🔄 Uploading..."):
                if uploaded_file.name.endswith(".csv"):
                    df = pd.read_csv(uploaded_file)
                    sheet_name = "CSV Data"
                else:
                    xls = pd.read_excel(uploaded_file, sheet_name=None)
                    sheet_names = list(xls.keys())
                    sheet_name = sheet_names[0]  # Default to first sheet
                    df = xls[sheet_name]
        
                st.session_state.data = df
                st.session_state.file_name = uploaded_file.name
                st.session_state.analysis_ready = True
                st.success(f"Loaded '{uploaded_file.name}' successfully.")
                st.rerun()
        except Exception as e:
                st.error(f"Failed to load file: {e}")
    elif not uploaded_file:
        st.error(f"Please upload your file")

def get_layout(tool):
    user_name = st.session_state.user_name.title()
    # logout_sidebar(user_name)
    output_dict = get_prompt(tool, user_name)
    query = st.chat_input(placeholder="Write your query?")
    if "messages" not in st.session_state:
        if tool not in ("📊 Data Analyzer Bot", "🔎 RAG-based Chatbot"):
            # st.title(output_dict['title'])
            st.markdown(output_dict['header'], unsafe_allow_html=True)
        st.session_state["messages"]=[]
        st.chat_message("assistant").write(output_dict['assistant_content'])
        
        if tool == "📊 Data Analyzer Bot":
            df = st.session_state.data
            st.chat_message("assistant").write("Here's a quick preview of your uploaded data:")
            st.chat_message("assistant").dataframe(df.head())    
    
    with st.expander("Previous Chat Messages"):
        for msg in st.session_state.messages:
            st.chat_message(msg["role"]).write(msg['content'])
    
    if user_name!='' and groq_api_key and query:
            
        st.session_state.messages.append({"role": "user", "content": query})
        st.chat_message("user").write(query)
        
        # Reconstruct chat history (excluding initial assistant greeting)
        chat_history = []
        for msg in st.session_state.messages[1:]:
            if msg["role"] == "user":
                chat_history.append(HumanMessage(content=msg["content"]))
            elif msg["role"] == "assistant":
                chat_history.append(AIMessage(content=msg["content"]))
    
        # Prompt template with username
        if tool == "📊 Data Analyzer Bot":
            df = st.session_state.data
            prompt_template = ChatPromptTemplate.from_messages([
            ("system", output_dict['system_prompt']),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}")
            ]).partial(username=user_name, query=query, columns=list(df.columns),
                        head=df.head().to_string(index=False))
    
        else:
            prompt_template = ChatPromptTemplate.from_messages([
                ("system", output_dict['system_prompt']),
                MessagesPlaceholder(variable_name="chat_history"),
                ("human", "{input}")
            ]).partial(username=user_name)
            
        if tool == "🔎 RAG-based Chatbot":
            hybrid_retriever = st.session_state.hybrid_retriever
            llm3 = ChatGroq(model="llama-3.3-70b-versatile",
               groq_api_key=groq_api_key,
                temperature = 0.7,  max_tokens = 400,   
                model_kwargs={"top_p" : 0.7,})
            contextualize_q_prompt = ChatPromptTemplate.from_messages(
                        [("system", contextualize_q_system_prompt),
                            MessagesPlaceholder("chat_history"),
                            ("human", "{input}"),])

            history_aware_retriever=create_history_aware_retriever(llm3, hybrid_retriever, contextualize_q_prompt)
            question_answer_chain = create_stuff_documents_chain(llm3, prompt_template)
            chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)
        else:
            llm3 = ChatGroq(model="llama-3.3-70b-versatile",
                           groq_api_key=groq_api_key,
                            temperature = 0.2,   max_tokens = 1000,   
                            model_kwargs={ "top_p" : 0.5, })
            chain: Runnable = prompt_template | llm3
        
        with st.chat_message("assistant"):
            st_cb=StreamlitCallbackHandler(st.container(),expand_new_thoughts=False)
            response=chain.invoke({"input": query,"chat_history": chat_history}, callbacks=[st_cb])
            if tool == "📊 Data Analyzer Bot":
                final_answer = response.content.strip('```python').strip("```").strip('python').strip('`')
                st.code(final_answer, language="python")
                # st.write(final_answer)
            elif tool == "🔎 RAG-based Chatbot":
                final_answer = response['answer']
                st.write(final_answer)
            else:
                final_answer = response.content if hasattr(response, "content") else str(response)
                st.write(final_answer)
            st.session_state.messages.append({'role': 'assistant', "content": final_answer})
        if tool == "📊 Data Analyzer Bot":
            # Safe execution (use caution in production)
            try:
                df_numeric = df.copy()
                local_vars = {'df': df_numeric}
                dependencies = {'pd': pd, 'px': px, 'np': np, 'scipy': scipy}
                exec(final_answer, dependencies, local_vars)
                result = local_vars.get('result', None)
                fig = local_vars.get("fig", None)
                if fig and hasattr(fig, 'to_plotly_json'):
                    key_hash = hashlib.md5(final_answer.encode()).hexdigest()
                    st.plotly_chart(fig, key=f"plot_{key_hash}")
            
                if result is not None:
                    st.write("### Result")
                    if isinstance(result, pd.DataFrame):
                        st.dataframe(result)
                    elif isinstance(result, str) and result.endswith(('.xlsx', '.csv')) and os.path.exists(result):
                        with open(result, 'rb') as f:
                            st.download_button("📥 Download Analysis File", f, file_name=os.path.basename(result))
                    else:
                        st.write(result)
            except IndexError as e:
                st.error(f"No matching records found. ({e})")
            except SyntaxError as e:
                st.error(f"Syntax error in generated code: {e}")
            except Exception as e:
                st.error(f"Error running code: {e}")
                    
    elif user_name!='' and groq_api_key and not query:
        if tool != "📊 Data Analyzer Bot":
            st.warning("Please type a query to get started.")

# Login
def login_screen():
    st.title("🔐 Login")
    with st.form("login_form", clear_on_submit=True):
        user_name = st.text_input("Enter your Name:")
        secret_key = st.text_input("Enter Secret Key:", type="password")
        submitted = st.form_submit_button("Login")
        if submitted:
            if secret_key == PASSWORD:
                st.session_state.logged_in = True
                st.session_state.user_name = user_name
                st.session_state.last_selected = "🏠 Home"
                st.rerun()  # To go straight to the main app
            elif user_name == '' or not user_name:
                st.error("Please enter your Name")
            elif not secret_key:
                st.error("Please enter Secret Key")
            else:
                st.error("Invalid Secret Key")

def home_screen():
    st.markdown(f"""
        <div style="padding: 15px; font-family: Arial, sans-serif;">
    <p style="font-size: 15px; line-height: 1.6;">
        Welcome to your <strong>Personal AI Assistant</strong> 👨‍💻, a versatile platform designed to empower users with a suite of intelligent tools that simplify complex tasks across various domains. Whether you are analyzing data, exploring documents, writing code, solving mathematical problems, or summarizing information — this app provides an intuitive, seamless experience tailored to your needs.
    </p>
    <h4 style="font-size: 17px;"><strong>Key Features:</strong></h4>
    <ul style="font-size: 15px; line-height: 1.8; padding-left: 15px;">
        <li><strong>📊 Excel Analyzer Bot</strong><br>
            Upload your datasets and interactively analyze them by simply asking questions in natural language. Get meaningful insights without writing a single line of code.
        </li>
        <li><strong>🔎 RAG-based Chatbot</strong><br>
            Upload documents in various formats and engage in intelligent, context-aware conversations. Ask detailed questions about your files and receive precise, AI-driven answers.
        </li>
        <li><strong>💻 Code Assistant</strong><br>
            Get instant help with programming questions across multiple languages. Whether it's debugging, code generation, or understanding complex concepts, your coding companion is here.
        </li>
        <li><strong>🧮 Math Assistant</strong><br>
            Solve mathematical problems effortlessly — from basic arithmetic to advanced equations. Simply state your math query and receive accurate solutions.
        </li>
        <li><strong>📝 Text Summarizer</strong><br>
            Submit lengthy text passages and receive concise, well-structured summaries, helping you grasp the key points faster and more efficiently.
        </li>
    </ul>
    <p style="font-size: 15px; line-height: 1.6;">
        This AI-powered application is crafted to <strong>boost productivity</strong>, <strong>enhance learning</strong>, and <strong>streamline problem-solving</strong> — making powerful technology accessible to everyone. Whether you're a student, a professional, or simply curious, this app is your all-in-one digital assistant ready to help you achieve more with ease.
    </p>
    <br>
<p style="text-align: center; font-size: 15px;">
    Developed with ❤️ by <strong>Rajkumar</strong>
</p>
</div>
    """, unsafe_allow_html=True)

# Streamlit UI                
def excel_analyser_screen(selection):
    get_layout(selection)
def RAG_based_chatbot_screen(selection):
    get_layout(selection)
def code_assistant_screen(selection):
    get_layout(selection)
def math_assistant_screen(selection):
    get_layout(selection)
def text_summarization_screen(selection):
    get_layout(selection)

def without_upload(selection):
    user = st.session_state.user_name.title()
    if selection == "💻 Code Assistant":
        code_assistant_screen(selection)
    elif selection == "🧮 Math Assistant":
        math_assistant_screen(selection)
    elif selection == "📝 Text Summarizer":
        text_summarization_screen(selection)

# Initialize session state
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if 'user_name' not in st.session_state:
    st.session_state.user_name = ''
if "last_selected" not in st.session_state:
    st.session_state.last_selected = "🏠 Home"
if 'analysis_ready' not in st.session_state:
    st.session_state.analysis_ready = False
if 'retriever_ready' not in st.session_state:
    st.session_state.retriever_ready = False


# Sidebar - Option Menu
with st.sidebar:
    st.markdown("""<style>
                section[data-testid="stSidebar"] {min-width: 350px; max-width: 350px; width: 350px;}
            </style>""", unsafe_allow_html=True)
    selected = option_menu(
        menu_title="Main Menu",
        options=["🏠 Home", "📊 Data Analyzer Bot", "🔎 RAG-based Chatbot", "💻 Code Assistant", "🧮 Math Assistant", "📝 Text Summarizer"],
        menu_icon="cast",
        default_index=0,)
    st.markdown("""
                <br>
                <center style="padding: 0px; margin: 0px;">
                    <a href="https://github.com/Rajiitb69/Gen_AI_streamlit_App/" target="_blank" style="margin-right: 10px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/733/733609.png" width="23">
                    </a>
                    <a href="mailto:raj345059@gmail.com" style="margin-right: 10px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/646/646094.png" width="27">
                    </a>
                    <a href="https://www.linkedin.com/in/raj-kumar-76b27993/" target="_blank">
                        <img src="https://cdn-icons-png.flaticon.com/512/174/174857.png" width="27">
                    </a>
                </center>
                """, unsafe_allow_html=True)

    

# Check if user changed selection
if selected != st.session_state.last_selected:
    st.session_state.last_selected = selected
    keys_to_keep = ["logged_in", "user_name","last_selected"]  # things you want to keep
    keys_to_delete = [key for key in st.session_state.keys() if key not in keys_to_keep]
    for key in keys_to_delete:
        del st.session_state[key]
    st.rerun()
    
# App Flow Control
selection = st.session_state.last_selected
if not st.session_state.logged_in:
    if selection == "🏠 Home":
        home_screen()
    else:
        login_screen()
        # st.warning("🔒 Please login to access other sections.")
else:
    if selection == "🏠 Home":
        home_screen()
    elif selection in ("💻 Code Assistant", "🧮 Math Assistant", "📝 Text Summarizer"):
        without_upload(selection)
    elif selection == "📊 Data Analyzer Bot":
        if not st.session_state.analysis_ready:
            data_analysis_uploader()
        else:
            excel_analyser_screen(selection)
    elif selection == "🔎 RAG-based Chatbot":
        if not st.session_state.retriever_ready:
            rag_chatbot_uploader()
        else:
            RAG_based_chatbot_screen(selection)

